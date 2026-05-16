"""
Report generation service.
"""
import os
import tempfile
from fpdf import FPDF
from docx import Document
from typing import Literal
from backend.models import DocumentFile, DocumentAnalysis
from .document_service import DocumentService

ReportFormat = Literal["markdown", "pdf", "docx"]


class ReportService:
    """Service for generating downloadable project reports."""

    @staticmethod
    def build_report_title(document: DocumentFile) -> str:
        return f"Relatório do documento: {document.filename}"

    @staticmethod
    def build_report_contents(document: DocumentFile, analysis: DocumentAnalysis) -> str:
        lines = [
            f"# {ReportService.build_report_title(document)}",
            "",
            f"**Documento:** {document.filename}",
            f"**Tipo:** {document.file_type}",
            f"**Tamanho:** {document.file_size} bytes",
            "",
            "## Sumário da Análise",
            analysis.summary or "Sem resumo disponível.",
            "",
            "## Observações de Organização",
            "A IA sugere revisar a estrutura do documento para clareza, consistência e priorização dos riscos identificados.",
            "",
            "## Riscos Identificados",
        ]

        if analysis.risks:
            for index, risk in enumerate(analysis.risks, start=1):
                lines.extend([
                    f"### Risco {index}",
                    f"- Descrição: {risk.description}",
                    f"- Severidade: {risk.severity or 'medium'}",
                    f"- Categoria: {risk.category or 'técnico'}",
                    f"- Mitigação: {risk.mitigation or 'Não informado'}",
                    "",
                ])
        else:
            lines.append("Nenhum risco identificado.")

        if analysis.full_analysis:
            lines.extend([
                "",
                "## Análise Completa",
                analysis.full_analysis,
            ])

        return "\n".join(lines)

    @staticmethod
    def generate_markdown_report(document: DocumentFile, analysis: DocumentAnalysis) -> str:
        return ReportService.build_report_contents(document, analysis)

    @staticmethod
    def generate_pdf_report(document: DocumentFile, analysis: DocumentAnalysis, output_path: str) -> str:
        content = ReportService.build_report_contents(document, analysis)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", "B", 16)
        pdf.multi_cell(0, 10, ReportService.build_report_title(document))
        pdf.ln(4)
        pdf.set_font("Arial", size=12)

        for line in content.split("\n"):
            pdf.multi_cell(0, 8, line)

        pdf.output(output_path)
        return output_path

    @staticmethod
    def generate_docx_report(document: DocumentFile, analysis: DocumentAnalysis, output_path: str) -> str:
        doc = Document()
        doc.add_heading(ReportService.build_report_title(document), level=1)
        doc.add_paragraph(f"Documento: {document.filename}")
        doc.add_paragraph(f"Tipo: {document.file_type}")
        doc.add_paragraph(f"Tamanho: {document.file_size} bytes")
        doc.add_paragraph("")
        doc.add_heading("Sumário da Análise", level=2)
        doc.add_paragraph(analysis.summary or "Sem resumo disponível.")
        doc.add_paragraph("")
        doc.add_heading("Riscos Identificados", level=2)

        if analysis.risks:
            for index, risk in enumerate(analysis.risks, start=1):
                doc.add_heading(f"Risco {index}", level=3)
                doc.add_paragraph(f"Descrição: {risk.description}")
                doc.add_paragraph(f"Severidade: {risk.severity or 'medium'}")
                doc.add_paragraph(f"Categoria: {risk.category or 'técnico'}")
                doc.add_paragraph(f"Mitigação: {risk.mitigation or 'Não informado'}")
        else:
            doc.add_paragraph("Nenhum risco identificado.")

        if analysis.full_analysis:
            doc.add_paragraph("")
            doc.add_heading("Análise Completa", level=2)
            doc.add_paragraph(analysis.full_analysis)

        doc.save(output_path)
        return output_path

    @staticmethod
    def create_report_file(document: DocumentFile, analysis: DocumentAnalysis, report_format: ReportFormat) -> str:
        supported = {"markdown": ".md", "pdf": ".pdf", "docx": ".docx"}
        if report_format not in supported:
            raise ValueError(f"Unsupported report format: {report_format}")

        suffix = supported[report_format]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            output_path = tmp_file.name

        if report_format == "markdown":
            content = ReportService.generate_markdown_report(document, analysis)
            with open(output_path, "w", encoding="utf-8") as file_obj:
                file_obj.write(content)
        elif report_format == "pdf":
            ReportService.generate_pdf_report(document, analysis, output_path)
        elif report_format == "docx":
            ReportService.generate_docx_report(document, analysis, output_path)

        return output_path
